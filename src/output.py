"""
output.py — write a translation result to a readable file (.txt, .docx, or .pdf).

The GUI lets a user pick an output format (or "same as input"). All three writers build
from the same structured sections, so the document always contains the plain-language
rewrite, the action list, the source quotes, the safety status, and the disclaimer.
"""

import os

SUPPORTED_FORMATS = ("txt", "docx", "pdf")


def resolve_format(choice: str, input_path: str) -> str:
    """Turn the user's choice (txt/docx/pdf/"same as input") into a concrete format."""
    choice = (choice or "txt").strip().lower()
    if choice in SUPPORTED_FORMATS:
        return choice
    if choice in ("same as input", "same", "match input"):
        ext = os.path.splitext(input_path)[1].lower().lstrip(".")
        return ext if ext in SUPPORTED_FORMATS else "txt"
    return "txt"


def _sections(result: dict) -> list[tuple[str, list[str]]]:
    """The document as (heading, lines) pairs — shared by all three writers."""
    conf = result.get("confidence", "unknown")
    if result.get("escalate"):
        safety = [
            f"Confidence: {conf}",
            "PLEASE SEE A PERSON: a caseworker or local legal aid should review this with you.",
            f"Why: {result.get('escalation_reason', '(not given)')}",
        ]
    else:
        safety = [f"Confidence: {conf}", "No escalation flagged."]

    return [
        ("Plain-language version", [result.get("plain_text", "(none returned)")]),
        ("What you need to do", list(result.get("action_items", []) or ["(none listed)"])),
        ("Where this comes from (quotes from your notice)",
         list(result.get("citations", []) or ["(none provided)"])),
        ("Confidence & safety", safety),
        ("Please note", [result.get("disclaimer", "")]),
    ]


def _write_txt(result: dict, path: str) -> None:
    blocks = []
    for heading, lines in _sections(result):
        blocks.append(heading.upper())
        blocks.append("-" * len(heading))
        blocks.extend(lines)
        blocks.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(blocks).strip() + "\n")


def _write_docx(result: dict, path: str) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Your notice, in plain language", level=0)
    for heading, lines in _sections(result):
        document.add_heading(heading, level=1)
        is_list = heading.startswith(("What you need", "Where this"))
        for line in lines:
            if not line:
                continue
            document.add_paragraph(line, style="List Bullet" if is_list else None)
    document.save(path)


def _write_pdf(result: dict, path: str) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    story = [Paragraph("Your notice, in plain language", styles["Title"]), Spacer(1, 12)]
    for heading, lines in _sections(result):
        story.append(Paragraph(heading, styles["Heading2"]))
        clean = [ln for ln in lines if ln]
        if heading.startswith(("What you need", "Where this")):
            items = [ListItem(Paragraph(ln, styles["BodyText"])) for ln in clean]
            story.append(ListFlowable(items, bulletType="bullet"))
        else:
            for ln in clean:
                story.append(Paragraph(ln.replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 10))
    SimpleDocTemplate(path, pagesize=letter).build(story)


def write_output(result: dict, fmt: str, path: str) -> str:
    """Write the result to `path` in format `fmt` (txt/docx/pdf). Returns the path."""
    writers = {"txt": _write_txt, "docx": _write_docx, "pdf": _write_pdf}
    if fmt not in writers:
        raise ValueError(f"Unsupported output format: {fmt}")
    writers[fmt](result, path)
    return path
